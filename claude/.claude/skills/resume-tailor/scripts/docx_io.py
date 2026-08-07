#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.10"
# dependencies = ["python-docx>=1.1"]
# ///
"""Read/write helper for resume .docx files.

extract <file.docx>
    Print paragraphs as JSON: [{"index", "style", "text", "runs":[{"text","bold","italic"}]}]

apply <master.docx> <out.docx> <edits.json>
    Clone master to out, applying edits. Ops (indexes refer to extract order of the master):
      {"op": "replace",      "index": N, "text": "..."}
      {"op": "delete",       "index": N}
      {"op": "insert_after", "index": N, "copy_format_from": M, "text": "..."}
      {"op": "replace_runs", "index": N, "runs": [{"text": "Label:", "bold": true},
                                                  {"text": "  rest", "bold": false}]}
    replace keeps the first run's formatting for the whole paragraph — don't use it
    on paragraphs with mixed bold/plain runs (check extract output first); use
    replace_runs there, which rebuilds the paragraph from explicit runs while
    inheriting the original first run's font.
"""
import json
import sys
from copy import deepcopy

from docx import Document


def extract(path):
    doc = Document(path)
    out = []
    for i, p in enumerate(doc.paragraphs):
        out.append({
            "index": i,
            "style": p.style.name,
            "text": p.text,
            "runs": [{"text": r.text, "bold": r.bold, "italic": r.italic} for r in p.runs],
        })
    json.dump(out, sys.stdout, indent=1, ensure_ascii=False)
    print()


def set_text(p, text):
    # Collapse to the first run (keeps its character formatting), drop the rest.
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r._element.getparent().remove(r._element)


def set_runs(p, runs_spec):
    # Rebuild from explicit runs; each clones the first run's rPr (font/size),
    # then overrides bold/italic as specified.
    from docx.text.run import Run
    template = deepcopy(p.runs[0]._element) if p.runs else None
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    for spec in runs_spec:
        if template is not None:
            el = deepcopy(template)
            p._element.append(el)
            run = Run(el, p)
            run.text = spec["text"]
        else:
            run = p.add_run(spec["text"])
        if "bold" in spec:
            run.bold = spec["bold"]
        if "italic" in spec:
            run.italic = spec["italic"]


def apply(master, out_path, edits_path):
    doc = Document(master)
    paras = list(doc.paragraphs)  # snapshot: original indexes stay valid across edits
    edits = json.load(open(edits_path))
    for e in edits:
        p = paras[e["index"]]
        if e["op"] == "replace":
            set_text(p, e["text"])
        elif e["op"] == "replace_runs":
            set_runs(p, e["runs"])
        elif e["op"] == "delete":
            p._element.getparent().remove(p._element)
        elif e["op"] == "insert_after":
            template = paras[e.get("copy_format_from", e["index"])]
            new = deepcopy(template._element)
            p._element.addnext(new)
            from docx.text.paragraph import Paragraph
            set_text(Paragraph(new, p._parent), e["text"])
        else:
            sys.exit(f"unknown op: {e['op']}")
    doc.save(out_path)
    print(f"wrote {out_path} ({len(edits)} edits)")


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "extract":
        extract(sys.argv[2])
    elif len(sys.argv) == 5 and sys.argv[1] == "apply":
        apply(sys.argv[2], sys.argv[3], sys.argv[4])
    else:
        sys.exit(__doc__)
